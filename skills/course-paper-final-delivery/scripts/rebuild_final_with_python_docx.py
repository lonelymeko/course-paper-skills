#!/usr/bin/env python3
"""Rebuild the final course paper using python-docx on top of a school template."""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path


EXTRA_DEPS = os.environ.get("COURSE_PAPER_PYTHON_DEPS")
if EXTRA_DEPS:
    sys.path.insert(0, EXTRA_DEPS)

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402


TITLE = "基于传感器、LoRaWAN与MQTT的温室智能灌溉物联网系统设计研究"
META = {
    "course": "物联网导论",
    "author": "张乐皓",
    "college": "软件学院",
    "class_name": "软件23-5",
    "instructor": "耿俊",
    "date": "2025 年 6 月 5 日",
}


def set_run_font(run, size_pt: float | None = None, name: str = "宋体", bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_body_format(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run, 12, "宋体")


def set_heading_format(paragraph, level: int) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(6 if level == 1 else 3)
    pf.space_after = Pt(6 if level == 1 else 3)
    size = 16 if level == 1 else 14 if level == 2 else 12
    for run in paragraph.runs:
        set_run_font(run, size, "黑体", True)


def set_caption_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, 10.5, "宋体")


def clear_document_after_cover(doc: Document, keep_paragraph_count: int | None = None) -> None:
    body = doc._body._element
    children = list(body)
    sect = None
    if children and children[-1].tag == qn("w:sectPr"):
        sect = children[-1]
        content_children = children[:-1]
    else:
        content_children = children

    if keep_paragraph_count is None:
        keep_paragraph_count = next(
            (idx + 1 for idx, para in enumerate(doc.paragraphs) if para.text.strip() == META["date"]),
            17,
        )

    for child in content_children[keep_paragraph_count:]:
        body.remove(child)

    if sect is not None and sect.getparent() is None:
        body.append(sect)


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def replace_paragraph_runs(paragraph, text: str, underline: bool = False, bold: bool = True, size: float = 16) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    set_run_font(run, size, "仿宋_GB2312", bold)
    run.underline = underline


def replace_text_runs_preserve_drawings(paragraph, text: str, size: float = 20) -> None:
    text_runs = [run for run in paragraph.runs if not run._r.xpath(".//w:drawing")]
    if text_runs:
        text_runs[0].text = text
        set_run_font(text_runs[0], size, "仿宋_GB2312", True)
        for run in text_runs[1:]:
            paragraph._p.remove(run._r)
        return

    for run in list(paragraph.runs):
        if run._r.xpath(".//w:drawing"):
            continue
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    set_run_font(run, size, "仿宋_GB2312", True)


def set_cover_paragraph_format(paragraph, line_pt: float = 24, after_pt: float = 0) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(0)
    pf.space_after = Pt(after_pt)


def fill_cover(doc: Document) -> None:
    fields = {
        7: ("课程名称：", META["course"], 12),
        9: ("作    者：", META["author"], 10),
        10: ("所在学院：", META["college"], 10),
        11: ("专业班级：", META["class_name"], 8),
        12: ("指导教师：", META["instructor"], 10),
    }
    for idx, (label, value, pad) in fields.items():
        p = doc.paragraphs[idx]
        for run in list(p.runs):
            p._p.remove(run._r)
        r1 = p.add_run(label)
        set_run_font(r1, 16, "仿宋_GB2312", True)
        r2 = p.add_run(value)
        set_run_font(r2, 16, "仿宋_GB2312", True)
        r2.underline = True
        if pad:
            r3 = p.add_run("\u00a0" * (pad * 3))
            set_run_font(r3, 16, "仿宋_GB2312", True)
            r3.underline = True
        set_cover_paragraph_format(p, line_pt=24, after_pt=0)

    title_p = doc.paragraphs[8]
    for run in list(title_p.runs):
        title_p._p.remove(run._r)
    title_label = title_p.add_run("题    目：")
    set_run_font(title_label, 16, "仿宋_GB2312", True)
    title_1 = title_p.add_run("基于传感器、LoRaWAN 与 MQTT")
    set_run_font(title_1, 16, "仿宋_GB2312", True)
    title_1.underline = True
    title_1.add_break()
    title_2 = title_p.add_run("的温室智能灌溉物联网系统设计研究")
    set_run_font(title_2, 16, "仿宋_GB2312", True)
    title_2.underline = True
    set_cover_paragraph_format(title_p, line_pt=24, after_pt=0)

    replace_paragraph_runs(doc.paragraphs[16], META["date"], underline=False, bold=True, size=16)
    doc.paragraphs[16].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cover_paragraph_format(doc.paragraphs[16], line_pt=24, after_pt=0)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def tighten_cover(doc: Document) -> None:
    # The template has several blank spacer paragraphs between the teacher field
    # and date. A long Chinese title wraps to two lines, so leaving every spacer
    # pushes the date to a separate page in LibreOffice/Word rendering.
    for idx in (15, 14, 13):
        if idx < len(doc.paragraphs) and not doc.paragraphs[idx].text.strip():
            remove_paragraph(doc.paragraphs[idx])


def iter_source_blocks(doc: Document):
    body = doc._body._element
    for child in body:
        if child.tag == qn("w:p"):
            para = next((p for p in doc.paragraphs if p._p is child), None)
            if para is not None:
                yield ("p", para)
        elif child.tag == qn("w:tbl"):
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is not None:
                yield ("tbl", table)


def copy_table(doc: Document, source_table) -> None:
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = source_table.cell(r, c).text.strip()
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 or c == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(18)
                for run in p.runs:
                    set_run_font(run, 10.5, "宋体", r == 0)


def append_source_content(target: Document, source: Document, fig_dir: Path) -> None:
    start = False
    table_idx = 0
    source_tables = list(source.tables)
    for kind, obj in iter_source_blocks(source):
        if kind == "p":
            text = obj.text.strip()
            if text == "摘 要":
                start = True
            if not start:
                continue
            if text in {"", "作者：待填写"}:
                continue
            if text.startswith("本文自绘的系统四层架构"):
                p = target.add_paragraph(text)
                set_paragraph_body_format(p)
                target.add_picture(str(fig_dir / "fig1-system-architecture.png"), width=Cm(13.5))
                target.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = target.add_paragraph("图1 温室智能灌溉系统四层架构图（自绘）")
                set_caption_format(cap)
                continue
            if text == "图1 温室智能灌溉系统四层架构图（自绘）":
                continue
            if text.startswith("监测流和控制流的关系"):
                p = target.add_paragraph(text)
                set_paragraph_body_format(p)
                target.add_picture(str(fig_dir / "fig2-data-control-flow.png"), width=Cm(13.5))
                target.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = target.add_paragraph("图2 数据流与灌溉控制闭环流程图（自绘）")
                set_caption_format(cap)
                continue
            if text == "图2 数据流与灌溉控制闭环流程图（自绘）":
                continue

            style = obj.style.name
            if style == "Heading 1":
                p = target.add_paragraph(text, style="Heading 1")
                set_heading_format(p, 1)
            elif style == "Heading 2":
                p = target.add_paragraph(text, style="Heading 2")
                set_heading_format(p, 2)
            elif style == "Heading 3":
                p = target.add_paragraph(text, style="Heading 3")
                set_heading_format(p, 3)
            elif re.match(r"表\d+\s", text):
                p = target.add_paragraph(text)
                set_caption_format(p)
            elif text == "新疆大学课程论文（设计）、学年论文评分表":
                p = target.add_paragraph(text, style="Heading 1")
                set_heading_format(p, 1)
            else:
                p = target.add_paragraph(text)
                set_paragraph_body_format(p)
        elif kind == "tbl":
            if not start:
                continue
            copy_table(target, source_tables[table_idx])
            table_idx += 1


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--fig-dir", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()

    doc = Document(args.template)
    source = Document(args.source)
    fill_cover(doc)
    tighten_cover(doc)
    clear_document_after_cover(doc)
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    append_source_content(doc, source, args.fig_dir)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)


if __name__ == "__main__":
    main()
